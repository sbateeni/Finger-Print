import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    doc.styles['Normal'].font.name = 'Traditional Arabic'
    doc.styles['Normal'].font.size = Pt(14)
    
    COLOR_PRIMARY = RGBColor(16, 44, 87)
    COLOR_SECONDARY = RGBColor(53, 114, 239)
    COLOR_TEXT = RGBColor(40, 40, 40)
    COLOR_HIGHLIGHT = RGBColor(194, 137, 34)
    
    # ---------------------------------------------------------------------------
    # Cover Page
    # ---------------------------------------------------------------------------
    for _ in range(5):
        doc.add_paragraph()
        
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("____________________________________________________")
    run_line.font.color.rgb = COLOR_HIGHLIGHT
    run_line.font.bold = True
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("تقرير تفصيلي وشامل حول مشروع\nنظام مطابقة البصمات الجنائي المدعوم بالرؤية الحاسوبية المتقدمة\n(Advanced CV Finger-Print System)")
    run_title.font.name = 'Traditional Arabic'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY
    
    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line2 = p_line2.add_run("____________________________________________________")
    run_line2.font.color.rgb = COLOR_HIGHLIGHT
    run_line2.font.bold = True
    
    for _ in range(3):
        doc.add_paragraph()
        
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run("إعداد: الضابط مقدم/ معتز عريدي\nنائب مدير مركز شرطة شقبا\nالتاريخ: يونيو 2026 م")
    run_meta.font.name = 'Traditional Arabic'
    run_meta.font.size = Pt(14)
    run_meta.font.bold = True
    run_meta.font.italic = True
    run_meta.font.color.rgb = COLOR_TEXT
    
    doc.add_page_break()
    
    # ---------------------------------------------------------------------------
    # Helper functions
    # ---------------------------------------------------------------------------
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Traditional Arabic'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(12)
        run_sub = p_sub.add_run("━" * 40)
        run_sub.font.color.rgb = COLOR_HIGHLIGHT
        run_sub.font.size = Pt(8)
        
    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Traditional Arabic'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        
    def add_body_paragraph(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.25
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            run_prefix = p.add_run(bold_prefix)
            run_prefix.font.name = 'Traditional Arabic'
            run_prefix.font.size = Pt(14)
            run_prefix.font.bold = True
            run_prefix.font.color.rgb = COLOR_PRIMARY
        run_text = p.add_run(text)
        run_text.font.name = 'Traditional Arabic'
        run_text.font.size = Pt(14)
        run_text.font.color.rgb = COLOR_TEXT
        return p

    def add_bullet_point(bold_title, description):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_title = p.add_run(bold_title + ": ")
        run_title.font.name = 'Traditional Arabic'
        run_title.font.size = Pt(13)
        run_title.font.bold = True
        run_title.font.color.rgb = COLOR_SECONDARY
        run_desc = p.add_run(description)
        run_desc.font.name = 'Traditional Arabic'
        run_desc.font.size = Pt(13)
        run_desc.font.color.rgb = COLOR_TEXT

    # ---------------------------------------------------------------------------
    # أولاً: مقدمة تمهيدية وتاريخية
    # ---------------------------------------------------------------------------
    add_heading_1("أولاً: مقدمة تمهيدية وتاريخية")
    add_body_paragraph("تعتبر البصمات من أهم وأقدم الأدلة الجنائية المادية التي يتم الاعتماد عليها في مسارح الجرائم للتعرف على الجناة وإثبات تواجدهم في مكان الحادث. يعود تاريخ استخدام البصمات إلى فترات طويلة، حيث أثبتت الدراسات العلمية والبيولوجية بما لا يدع مجالاً للشك أن بصمة الإصبع هي سمة فريدة لا يمكن أن تتطابق بين شخصين على وجه الأرض، ولا حتى بين التوائم المتطابقة. ومع التطور السريع للتكنولوجيا وثورة المعلومات، انتقلت الأجهزة الأمنية والشرطية من المطابقة اليدوية التقليدية التي كانت تستغرق أسابيع وشهوراً وتتطلب مجهوداً بشرياً هائلاً، إلى الاعتماد على الأنظمة الآلية للتعرف على البصمات (AFIS).")
    add_body_paragraph("في هذا السياق، يأتي مشروعنا \"نظام مطابقة البصمات الجنائي المدعوم بالرؤية الحاسوبية المتقدمة (Finger-Print)\" كخطوة متقدمة ونقلة نوعية تتجاوز الأنظمة التقليدية، حيث يدمج أحدث تقنيات الرؤية الحاسوبية (Computer Vision) والخوارزميات الرياضية الدقيقة لتوفير منصة موثوقة، دقيقة، وسريعة. تم تصميم هذا النظام ليكون أداة مساعدة (Decision Support System) لخبراء الأدلة الجنائية، مما يسهم بشكل مباشر في تسريع وتيرة التحقيقات الجنائية، وتقليل تراكم القضايا، وتعزيز منظومة العدالة الجنائية بقرائن وأدلة قطعية تواكب التحديات المعاصرة.")

    # ---------------------------------------------------------------------------
    # ثانياً: الآلية المعمارية والتقنية
    # ---------------------------------------------------------------------------
    add_heading_1("ثانياً: الآلية المعمارية والتقنية (ما يقوم به النظام بالتفصيل)")
    add_body_paragraph("يعمل النظام من خلال سلسلة مترابطة من الخوارزميات التي تحاكي طريقة عمل الخبير البشري ولكن بسرعة فائقة وقدرة تحليلية أدق. يمكن تقسيم مهام النظام إلى المراحل التقنية التالية:")

    add_heading_2("1. معالجة الصور وتحسينها الاستباقي (Image Preprocessing and Enhancement)")
    add_body_paragraph("غالباً ما تكون البصمات المرفوعة من مسارح الجرائم غير واضحة، أو مشوهة، أو متداخلة مع خلفيات معقدة. يقوم النظام بعملية تحسين آلية للصور تشمل:")
    add_bullet_point("إزالة الضوضاء والتشويش", "باستخدام مرشح CLAHE (Contrast Limited Adaptive Histogram Equalization) لتحسين التباين وعزل الخطوط الأساسية للبصمة عن الخلفية البيئية المحيطة، مما يسهل دقة القراءة.")
    add_bullet_point("تطبيق فلاتر غابور (Gabor Filters)", "وهي فلاتر رياضية متقدمة تقوم بتقوية وتوضيح خطوط البصمة الباهتة بناءً على اتجاهاتها وتردداتها الأصلية بدقة متناهية، مع إعدادات قابلة للتعديل (حجم الكتلة = 11، الثابت C = 2).")
    add_bullet_point("التحويل الثنائي والتنحيف (Binarization & Thinning)", "تحويل صورة البصمة إلى هيكل عظمي أحادي البكسل باستخدام خوارزمية (Zhang-Suen)، مما يسهل قراءة التفاصيل الدقيقة دون التأثر بمدى قوة ضغط الإصبع.")

    add_heading_2("2. استخراج النقاط الدقيقة وتحليلها (Minutiae Points Extraction)")
    add_body_paragraph("النقاط الدقيقة هي العلامات المميزة في أي بصمة وتعتمد عليها المحاكم دولياً كدليل قطعي صلب. يقوم النظام تلقائياً بمسح الهيكل العظمي وتحليل نمط الجيران الثمانية (8-Neighbor) لكل بكسل لتحديد العناصر الجوهرية الآتية:")
    add_bullet_point("نهايات الخطوط (Ridge Endings)", "تحديد الإحداثيات الدقيقة للنقطة التي ينتهي عندها خط البصمة فجأة بصورة هندسية واضحة.")
    add_bullet_point("التفرعات (Bifurcations)", "رصد النقطة الدقيقة التي ينقسم فيها خط واحد رئيسي من خطوط البصمة إلى خطين فرعيين.")
    add_bullet_point("النقاط والجزر (Dots and Islands)", "تحليل الخطوط القصيرة جداً أو النقاط المعزولة الواقعة في الفراغات البينية للخطوط الأخرى.")
    add_body_paragraph("يقوم النظام بتسجيل الإحداثيات الهندسية المعتمدة على محاور الديكارتي (X, Y) لكل نقطة مرصودة، بالإضافة إلى زاوية الميلان والاتجاه، ثم يخزنها في قاعدة البيانات كـ \"توقيع رياضي\" مشفر وفريد للبصمة.")

    add_heading_2("3. تحليل الأنماط الكلية للخطوط (Global Ridge Patterns)")
    add_body_paragraph("إلى جانب النقاط الدقيقة، يقوم النظام بتصنيف البصمة بناءً على شكلها ونمطها الهيكلي الكلي (النوع، المنطقة، النمط) لتضييق نطاق البحث. يسهم هذا التصنيف الأولي في تقليص وقت البحث من ملايين السجلات إلى بضعة آلاف فقط في ثوان معدودة.")

    add_heading_2("4. خوارزميات المطابقة الدقيقة والمطابقة الجزئية")
    add_body_paragraph("تعتبر هذه الميزة الأهم والأبرز في النظام، نظراً لأن البصمات المرفوعة جنائياً نادراً ما تكون كاملة بنسبة 100%. الخوارزميات الرياضية في هذا النظام قادرة على تقديم كفاءة استثنائية من خلال:")
    add_bullet_point("المطابقة الجزئية الذكية", "مقارنة النقاط المستخرجة من البصمة الجزئية مع النقاط المقابلة في البصمة الأصلية باستخدام معايير المسافة المكانية (Euclidean Distance) والفرق الزاوي (Angular Deviation)، وتكوين أزواج مطابقة محتملة.")
    add_bullet_point("التكيف مع الدوران والإزاحة", "القدرة التلقائية على تدوير التوقيع الرياضي بجميع الزوايا الممكنة لإيجاد التطابق بغض النظر عن زاوية طبع البصمة وقت الحادث.")
    add_bullet_point("حساب درجة الثقة (Confidence Scoring)", "حساب مقياس الثقة بناءً على نسبة النقاط المتطابقة إلى إجمالي النقاط في البصمة الجزئية، مع تصنيف النتيجة وفق مقاييس NIST (دقيقة/متوسطة/ضعيفة/غير متطابقة).")

    add_heading_2("5. إدارة النتائج وواجهة المراجعة البشرية (Manual Editor)")
    add_body_paragraph("يولد النظام تقارير فنية مفصلة توضح نسبة التطابق الرقمي بين البصمة المجهولة والبصمات المشتبه بها، مع رسم خطوط توصيل بصرية ملونة بين النقاط المتطابقة على الشاشة. يقدم النظام هذه النتيجة التلقائية عبر واجهة مستخدم مخصصة (محرر يدوي تفاعلي) تتيح للخبير الجنائي:")
    add_bullet_point("عرض جانبي مزدوج", "عرض البصمة الأصلية والجزئية جنباً إلى جنب بنفس المقياس للمقارنة البصرية المباشرة.")
    add_bullet_point("إضافة وحذف وتعديل النقاط", "النقر للإضافة في وضع الإضافة، الحذف المباشر في وضع الحذف، وتغيير النوع التشريحي عبر واجهة منسدلة.")
    add_bullet_point("سحب نقاط المطابقة وتحريكها", "تصحيح أزواج المطابقة يدوياً عن طريق سحب النقاط على طبقة التطابق (Match Overlay) أو تحريك أي نقطة إلى موقعها الصحيح.")
    add_bullet_point("تصدير التقرير", "اعتماد التعديلات وإنشاء تقرير مطابقة نهائي بصيغة DOCX مع التوثيق الرقمي الكامل للإجراءات المتخذة (Audit Trail) لضمان النزاهة القانونية والشفافية التامة أمام القضاء.")

    # ---------------------------------------------------------------------------
    # ثالثاً: المزايا الاستراتيجية والتشغيلية
    # ---------------------------------------------------------------------------
    add_heading_1("ثالثاً: المزايا الاستراتيجية والتشغيلية (إيجابيات النظام)")
    add_body_paragraph("من شأن تبني هذا النظام المتطور أن ينعكس إيجاباً وبقوة على كفاءة وجودة أداء مختبرات البحث الجنائي من خلال عدة مميزات رائدة:")

    add_bullet_point("1. السرعة القياسية وتقليص التراكمات", "بدلاً من استغراق أيام لمقارنة بصمة واحدة بآلاف السجلات يدوياً، يمكن للنظام تحليل ومقارنة البصمات في ثوان معدودة، مما يسهم في تسريع القضايا المعلقة.")
    add_bullet_point("2. دقة متناهية تفوق القدرات البشرية", "قدرة الخوارزميات الحاسوبية على رصد التفاصيل فائقة الصغر يقلل من نسبة النتائج السلبية الخاطئة الناتجة عن تعب وإجهاد العين البشرية للخبير الجنائي.")
    add_bullet_point("3. حل معضلة البصمات الجزئية", "يستطيع النظام فك شفرة البصمات غير الكاملة التي كانت تقيد قضاياها سابقاً ضد مجهول، مما يمنح جهات التحقيق دليلاً مادياً إضافياً.")
    add_bullet_point("4. تخفيف العبء الإدراكي عن الخبراء", "يقوم النظام بأعمال الفلترة الأولية ويقترح قائمة بالمطابقات المحتملة، لتوجيه جهود الخبير نحو التحقق والاعتماد النهائي للتقرير.")
    add_bullet_point("5. الشفافية الكاملة (Explainability)", "عكس أنظمة 'الصندوق الأسود' للذكاء الاصطناعي، يعتمد النظام على خوارزميات رياضية مفتوحة المصدر وقابلة للتفسير العلمي أمام القاضي، مما يعزز موقف الخبير الجنائي قانونياً.")
    add_bullet_point("6. القدرة على التوسع (Scalability)", "تصميم مرن يتيح دمج النظام مع قواعد بيانات السجل المدني الوطنية عبر واجهات برمجية آمنة (APIs).")

    # ---------------------------------------------------------------------------
    # رابعاً: التحديات التقنية والتشغيلية
    # ---------------------------------------------------------------------------
    add_heading_1("رابعاً: التحديات التقنية والتشغيلية (المعيقات وما يجب مراعاته)")
    add_body_paragraph("لضمان الشفافية والموضوعية التامة، ورغم كل الإمكانيات المتقدمة التي يوفرها النظام، فإنه يواجه بعض التحديات التشغيلية والمحددات الفنية التي تستدعي تخطيطاً حذراً:")

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    shading_elm = parse_xml(r'<w:shd {} w:fill="F4F4F6"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'''
        <w:tcBorders {}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="single" w:sz="36" w:space="0" w:color="C28922"/>
        </w:tcBorders>
    '''.format(nsdecls('w')))
    tcPr.append(tcBorders)
    
    p_box = cell.paragraphs[0]
    p_box.paragraph_format.space_before = Pt(6)
    p_box.paragraph_format.space_after = Pt(6)
    
    r_box_title = p_box.add_run("⚠️ تنويه فني وتحديات استراتيجية هامة للتشغيل:\n\n")
    r_box_title.font.name = 'Traditional Arabic'
    r_box_title.font.size = Pt(13)
    r_box_title.font.bold = True
    r_box_title.font.color.rgb = COLOR_HIGHLIGHT
    
    r_box_desc = p_box.add_run(
        "1. حتمية جودة المدخلات (Garbage In, Garbage Out): إذا كانت البصمة مشوهة وتالفة بالكامل دون نقاط مرجعية كافية، فلن يقدم النظام نتائج موثوقة وجاري العمل على حل هذه المشكلة.\n"
        "2. موثوقية الأدلة أمام القضاء: يعتمد النظام على خوارزميات رياضية شفافة وقابلة للتفسير العلمي أمام القاضي، مما يعزز موقف الخبير الجنائي قانونياً.\n"
        "3. التشوهات المرنة (Elastic Distortion): اختلاف ضغط وسحب البصمة عند التسجيل مقابل المسرح يؤثر على دقة المطابقة ويتطلب خوارزميات متحملة للتشوه.\n"
        "4. أمن وحساسية البيانات البيومترية: ضرورة تشغيل النظام على شبكات داخلية آمنة مع تشفير قوي لحماية خصوصية وسرية البيانات البيومترية للأفراد."
    )
    r_box_desc.font.name = 'Traditional Arabic'
    r_box_desc.font.size = Pt(12)
    r_box_desc.font.color.rgb = COLOR_TEXT
    
    doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # خامساً: الخلاصة والتوصيات الختامية
    # ---------------------------------------------------------------------------
    add_heading_1("خامساً: الخلاصة والتوصيات الختامية")
    add_body_paragraph("يمثل مشروع \"نظام مطابقة البصمات الجنائي المدعوم بالرؤية الحاسوبية المتقدمة (Finger-Print)\" استثماراً استراتيجياً حقيقياً لأي مؤسسة أمنية تسعى لتطوير أدواتها ومواكبة الجريمة الحديثة في العصر الرقمي. إنه لا يستبدل الخبير الجنائي البشري، بل يمنحه أداة تضاعف من سرعته وإنتاجيته وموثوقيته عشرات المرات. من خلال معالجة الصور، واستخراج النقاط الدقيقة، وإدارة المطابقات الجزئية بفعالية، يمكن للنظام أن يفك طلاسم العديد من الجرائم الغامضة ويدعم أركان العدالة بكفاءة واقتدار.")
    add_body_paragraph("ولتحقيق الاستفادة القصوى والنجاح المستدام لهذا المشروع، نوصي بتبني النقاط التالية:")
    add_bullet_point("توفير خوادم عالية الأداء", "تأسيس بنية تحتية مخصصة لضمان تلبية طلبات البحث المتزامنة دون تأخير في أداء المطابقة.")
    add_bullet_point("تخصيص دورات تدريبية متقدمة", "تأهيل الكوادر الشرطية ومحققي مسرح الجريمة على الاستخدام الأمثل للمحرر اليدوي وحالات البصمات الجزئية.")
    add_bullet_point("تحسين خوارزميات المطابقة", "تطوير خوارزمية مطابقة أكثر تحملاً للتشوهات (Distortion-Tolerant Matching) باستخدام تقنيات المطابقة المرنة (Elastic Matching).")
    add_bullet_point("قاعدة بيانات مركزية للبصمات", "ربط النظام بقاعدة بيانات مركزية (AFIS) للمقارنة الآلية على نطاق واسع مع قواعد بيانات الإنتربول والسجل المدني.")
    add_bullet_point("توثيق واجهات API", "توثيق كامل لواجهات REST API لتسهيل الدمج مع الأنظمة الجنائية الأخرى.")
    add_bullet_point("إصدار جوال للميدان", "تطبيق جوال لالتقاط البصمات وإرسالها للتحليل الفوري من موقع الحدث.")

    # ---------------------------------------------------------------------------
    # تذييل: حالة التطوير
    # ---------------------------------------------------------------------------
    add_heading_1("سادساً: حالة التطوير والختام")
    add_body_paragraph("يُشير هذا التقرير إلى أن مشروع \"Finger-Print\" لا يزال في مرحلة التطوير النشط (Active Development). تم إنجاز مراحل معالجة الصور، استخراج النقاط الدقيقة، المطابقة الجزئية، والمحرر اليدوي التفاعلي. لا يزال العمل جارياً على تحسين أداء المطابقة وتوسيع قدرات النظام.")

    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run("— — — — — — — — — — — — — — — — — — — — — — — — —")
    run_footer.font.color.rgb = COLOR_HIGHLIGHT
    run_footer.font.size = Pt(10)

    p_dev = doc.add_paragraph()
    p_dev.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_dev = p_dev.add_run("تنبيه: هذا البرنامج لا يزال قيد التطوير (Under Development)\nقد تطرأ تغييرات على الوظائف والواجهات في الإصدارات القادمة")
    run_dev.font.name = 'Traditional Arabic'
    run_dev.font.size = Pt(12)
    run_dev.font.bold = True
    run_dev.font.color.rgb = COLOR_HIGHLIGHT

    p_footer2 = doc.add_paragraph()
    p_footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer2 = p_footer2.add_run("— — — — — — — — — — — — — — — — — — — — — — — — —")
    run_footer2.font.color.rgb = COLOR_HIGHLIGHT
    run_footer2.font.size = Pt(10)

    # RTL
    for p in doc.paragraphs:
        pPr = p._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
        for run in p.runs:
            rPr = run._element.get_or_add_rPr()
            rtl = OxmlElement('w:rtl')
            rtl.set(qn('w:val'), '1')
            rPr.append(rtl)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pPr = p._element.get_or_add_pPr()
                    bidi = OxmlElement('w:bidi')
                    bidi.set(qn('w:val'), '1')
                    pPr.append(bidi)
                    for run in p.runs:
                        rPr = run._element.get_or_add_rPr()
                        rtl = OxmlElement('w:rtl')
                        rtl.set(qn('w:val'), '1')
                        rPr.append(rtl)

    return doc

doc = create_document()
doc.save("Finger_Print_System_Report.docx")
